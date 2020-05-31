import pandas as pd
import io
import csv
import docx
import datetime as datetime
import numpy as np
from git import Repo

date = datetime.datetime.now().strftime('%Y-%m-%d')

def read_docx_tables(doc, tab_id=None, **kwargs):

    def read_docx_tab(tab, **kwargs):
        vf = io.StringIO()
        writer = csv.writer(vf)
        for row in tab.rows:
            writer.writerow(cell.text for cell in row.cells)
        vf.seek(0)
        if len(tab.rows) > 0:
            return pd.read_csv(vf, **kwargs)
        else:
            return pd.DataFrame()

    if tab_id is None:
        return [read_docx_tab(tab, **kwargs) for tab in doc.tables]
    else:
        return read_docx_tab(doc.tables[tab_id], **kwargs)


def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)


def log_table(doc, table_name, comments=[]):

    if len(comments) == 0:
        table = doc.add_table(0, 0)

    else:
        if table_name == "goals":
            doc.add_heading("Goals")
            table = doc.add_table(1, 3)
            heading_cells = table.rows[0].cells
            heading_cells[0].text = "Project"
            heading_cells[1].text = "Goal"
            heading_cells[2].text = "Summary"

        if table_name == "accomplished":
            doc.add_heading("Accomplished")
            table = doc.add_table(1, 3)
            heading_cells = table.rows[0].cells
            heading_cells[0].text = "Project"
            heading_cells[1].text = "Accomplished"
            heading_cells[2].text = "Notes"

        if table_name == "pushes":
            doc.add_heading("Code Executions")
            table = doc.add_table(1, 3)
            heading_cells = table.rows[0].cells
            heading_cells[0].text = "Project"
            heading_cells[1].text = "URL"
            heading_cells[2].text = "Comment"

        if table_name == "to_do":
            doc.add_heading("To-Do List")
            table = doc.add_table(1, 3)
            heading_cells = table.rows[0].cells
            heading_cells[0].text = "Project"
            heading_cells[1].text = "To Do"
            heading_cells[2].text = "Notes"

        for row in comments:
            cells = table.add_row().cells
            cells[0].text = row[0]
            cells[1].text = row[1]
            cells[2].text = row[2]

        table.style = 'ColorfulList-Accent2'

def new_entry(doc, notes):
    heading = doc.add_paragraph(date)

    goals = notes['goals']
    accomplished = notes['accomplished']
    pushes = notes['pushes']
    to_do = notes['to_do']
    log_table(doc, "goals", goals)
    log_table(doc, "accomplished", accomplished)
    log_table(doc, "pushes", pushes)
    log_table(doc, "to_do", to_do)


def table_merge(doc, table, comments, name, index):
    if len(table) == 0 and len(comments) > 0:
        log_table(doc, name, comments)

    elif len(table) == 0 and len(comments) == 0:
        return

    else:
        if name == 'goals':
            columns = ["Project", "Goal", "Summary"]
        elif name == 'accomplished':
            columns = ["Project", "Accomplished", "Notes"]
        elif name == 'pushes':
            columns = ["Project", "URL", "Comment"]
        elif name == 'to_do':
            columns = ["Project", "To Do", "Notes"]

        update = pd.DataFrame(comments, columns=columns)

        new = pd.merge(table.replace(np.NaN, ''), update,
                       'outer').drop_duplicates(subset=columns[1],keep='last').values

        current_table = doc.tables[index]
        current_length = len(current_table.rows) - 2
        for i, row in enumerate(new):
            if i <= current_length:
                current_table.row_cells(i+1)[0].text = new[i][0]
                current_table.row_cells(i+1)[1].text = new[i][1]
                current_table.row_cells(i+1)[2].text = new[i][2]
            elif i > current_length:
                cells = current_table.add_row().cells
                cells[0].text = new[i][0]
                cells[1].text = new[i][1]
                cells[2].text = new[i][2]

def update_entry(doc, notes):
    goals = notes['goals']
    accomplished = notes['accomplished']
    pushes = notes['pushes']
    to_do = notes['to_do']

    for i, table in enumerate(reversed(doc.tables)):
        ind = len(doc.tables)-1-i
        table = read_docx_tables(doc, ind)
        if len(table) > 0:
            table_title = table.columns[1]
        else:
            table_title = "NA"
        if table_title == 'Goal':
            goal_ind = ind
            goal_tab = table
            table_merge(doc, goal_tab, goals, 'goals', goal_ind)
            break

    for i, table in enumerate(doc.tables[goal_ind:]):
        i = goal_ind+i
        table = read_docx_tables(doc, i)

        if len(table) > 0:
            table_title = table.columns[1]
        else:
            table_title = "NA"

        if table_title == 'Accomplished':
            acc_ind = i
            acc_tab = table
            table_merge(doc, acc_tab, accomplished, 'accomplished', acc_ind)

        if table_title == 'URL':
            push_ind = i
            push_tab = table
            table_merge(doc, push_tab, pushes, 'pushes', push_ind)

        if table_title == 'To Do':
            td_ind = i
            td_tab = table
            table_merge(doc, td_tab, to_do, 'to_do', td_ind)


def git_push(project, PATH_OF_GIT_REPO):
    COMMIT_MESSAGE = input("Commit Message");
    try:
        repo = Repo(PATH_OF_GIT_REPO)
        repo.git.add('-A')
        repo.index.commit(COMMIT_MESSAGE)
        origin = repo.remote(name='origin')
        origin.push()
        sha = repo.head.object.hexsha
        remote_url = repo.remotes[0].config_reader.get("url")
        repo_url = remote_url.split(":", 1)[1].split(".git", 1)[0]
        repo_url = "https://github.com/"+repo_url+"/commit/"+sha

        return [[project, repo_url, COMMIT_MESSAGE]]

    except:
        print('Some error occured while pushing the code')


def validate(date_text):
    try:
        datetime.datetime.strptime(date_text, '%Y-%m-%d')
        return True
    except ValueError:
        return False

def push_update(doc_path, PATH_OF_GIT_REPO, notes):
    doc=docx.Document(doc_path)
    project=notes['project']
    if PATH_OF_GIT_REPO is not None:
        notes.update({'pushes':git_push(project,PATH_OF_GIT_REPO)})
    else:
        notes.update({'pushes':[]})
    for i, text in enumerate(reversed(doc.paragraphs)):
        if validate(text.text) == True:
            last_date = text.text
            index=len(doc.paragraphs)-1-i
            break
    if last_date != date:
        new_entry(doc, notes)
        doc.save(doc_path)
        return("New Entry Added")
    if last_date == date:
        update_entry(doc, notes)
        doc.save(doc_path)
        return("Updated Entry")
